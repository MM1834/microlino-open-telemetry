from pathlib import Path

from docx import Document


SOURCE = Path("/Users/martin/Downloads/MOT_Remote_Pilot_Onboarding_Erweitert.docx")
OUTPUT = Path("/Users/martin/Documents/MICROLINO/microlino-open-telemetry/output/docx/MOT_Remote_Pilot_Onboarding_Extended_EN.docx")


TRANSLATIONS = {
    "MOT REMOTE-PILOT ONBOARDING": "MOT REMOTE PILOT ONBOARDING",
    "MOT Remote-Pilot Onboarding | Erweiterte Arbeitsversion": "MOT Remote Pilot Onboarding | Extended working version",
    "Seite": "Page",
    "Erweiterte Arbeitsanleitung": "Extended Setup Guide",
    "Portal-Account, lokale Adapter-Einrichtung und Claim": "Portal account, local adapter setup and vehicle claim",
    "Dokumentstruktur": "Document structure",
    "1  MOT-Portal-Konto aktivieren": "1  Activate the MOT Portal account",
    "2  Adapter vollständig lokal einrichten": "2  Complete the local adapter setup",
    "3  Adapter und Portal über den Claim verbinden": "3  Link the adapter and portal using the claim",
    "Stand: 28.08.2026": "Version: 28 August 2026",
    "MOT-Portal-Konto aktivieren": "Activate the MOT Portal account",
    "Dieser Teil kann bereits erledigt werden, bevor der Adapter beim Pilotbenutzer eintrifft.": "This phase can be completed before the adapter arrives with the pilot user.",
    "1  Einladung erhalten": "1  Receive the invitation",
    "E-Mail mit Betreff MOT Portal - Ihr Zugang öffnen. Falls sie fehlt: Spam-Ordner prüfen.": "Open the email with the subject MOT Portal - Your access. If it is missing, check the spam folder.",
    "Abbildung 1.1 - Screenshot ergänzen": "Figure 1.1 - Add screenshot",
    "2  Portal öffnen": "2  Open the portal",
    "Abbildung 1.2 - Screenshot ergänzen": "Figure 1.2 - Add screenshot",
    "3  Erstmals anmelden": "3  Sign in for the first time",
    "Benutzername und temporäres Passwort aus der Einladung verwenden.": "Use the username and temporary password from the invitation.",
    "Abbildung 1.3 - Screenshot ergänzen": "Figure 1.3 - Add screenshot",
    "4  Persönliches Passwort setzen": "4  Set your personal password",
    "Ein eigenes sicheres Passwort festlegen und sicher aufbewahren.": "Choose your own secure password and store it safely.",
    "Abbildung 1.4 - Screenshot ergänzen": "Figure 1.4 - Add screenshot",
    "5  Account bereit": "5  Account ready",
    "Der Status ist nun CONFIRMED. Noch kein Fahrzeug sichtbar ist zu diesem Zeitpunkt normal.": "The status is now CONFIRMED. It is normal that no vehicle is visible at this stage.",
    "Abbildung 1.5 - Screenshot ergänzen": "Figure 1.5 - Add screenshot",
    "Adapter lokal einrichten": "Set up the adapter locally",
    "Alle Schritte dieser Phase erfolgen direkt am Adapter mit Laptop, Smartphone oder Tablet.": "Complete every step in this phase directly on the adapter using a laptop, smartphone or tablet.",
    "Hinweis: \nDer Adapter bracht immer Strom über den USB-C Anschluss. Es wird aktuell kein Strom über den OBD2 bezogen.": "Note: The adapter always requires power through its USB-C port. It currently receives no power through OBD-II.",
    "1  Adapter einschalten": "1  Power on the adapter",
    "Adapter mit Strom versorgen. Für die Ersteinrichtung muss er noch nicht am Fahrzeug sein.": "Connect the adapter to USB-C power. It does not need to be connected to the vehicle for the initial setup.",
    "Abbildung 2.1 - Adapter mit USB verbinden": "Figure 2.1 - Connect the adapter to USB-C power",
    "2  Mit dem MOT-WLAN verbinden": "2  Connect to the MOT Wi-Fi network",
    "WLAN/WiFi Access Point MOT-xxxx auswählen. Das Passwort gemäss Inventarblatt verwenden.": "Select the MOT-xxxx Wi-Fi access point. Use the password shown on the inventory sheet.",
    "Abbildung 2.2 - Mit WLAN verbinden": "Figure 2.2 - Connect to Wi-Fi",
    "3  Lokale Setup-Seite öffnen": "3  Open the local setup page",
    "Im Browser http://192.168.4.1/ öffnen und als setup mit dem Passwort gemäss Inventarblatt anmelden. Es ist dasselbe Passwort wie für das WLAN.": "Open http://192.168.4.1/ in a browser and sign in as setup using the password on the inventory sheet. This is the same password as for the Wi-Fi network.",
    "Abbildung 2.3 - mit dem Browser zum MOT Device verbinden": "Figure 2.3 - Connect to the MOT device in a browser",
    "4  Neues Passwort setzen": "4  Set a new password",
    "Das neue Admin-/WLAN-Passwort zweimal identisch eingeben, speichern und das Passwort merken.": "Enter the new admin/Wi-Fi password twice, save it and remember the password.",
    "Abbildung 2.4 - Passwort setzen": "Figure 2.4 - Set the password",
    "5  Nach Neustart neu verbinden": "5  Reconnect after the restart",
    "Erneut mit MOT-xxxx verbinden - jetzt mit dem neuen Passwort. Danach http://192.168.4.1/ öffnen und unbedingt mit Benutzername admin anmelden, nicht setup.": "Reconnect to MOT-xxxx using the new password. Then open http://192.168.4.1/ and make sure you sign in with the admin username, not setup.",
    "Abbildung 2.5.1 - Neues WLAN Passwort gesetzt": "Figure 2.5.1 - New Wi-Fi password saved",
    "Abbildung 2.5.2 - WLAN - Neues Passwort": "Figure 2.5.2 - Wi-Fi - New password",
    "6  Connectivity im Wizard konfigurieren": "6  Configure connectivity in the wizard",
    "Im Wizard jeweils mit Next oder Save... weitergehen. Home- oder Mobile/WiFi2 eintragen. Nur 2.4 GHz; beim iPhone Kompatibilität maximieren aktivieren.": "Use Next or Save... to move through the wizard. Enter a Home Wi-Fi or Mobile/WiFi2 network. Use 2.4 GHz only; on an iPhone, enable Maximise Compatibility for the Personal Hotspot.",
    "Abbildung 2.6 - Als admin einloggen, mit neuem Passwort": "Figure 2.6 - Sign in as admin using the new password",
    "7  WLAN speichern und neu verbinden": "7  Save the Wi-Fi settings and reconnect",
    "Es können bis zwei WLAN konfiguriert werden. Das HOME WLAN wird hat dabei höhere Priorität und wird immer genutzt wenn es verfügbar ist. Mobile WLAN wird aktiv wenn das HOME WLAN nicht verfügbar ist. Bei einem WLAN kann diese als HOME oder Mobil eingetragen werden.": "You can configure up to two Wi-Fi networks. Home Wi-Fi has priority and is used whenever it is available. Mobile Wi-Fi becomes active when Home Wi-Fi is unavailable. If you use only one network, enter it as either Home or Mobile Wi-Fi.",
    "Bei geänderten WLAN-Daten startet der Adapter neu. Danach die Verbindung neu herstellen und wieder http://192.168.4.1/ öffnen.": "The adapter restarts after Wi-Fi settings are changed. Reconnect afterwards and open http://192.168.4.1/ again.",
    "Abbildung 2.7 - Screenshot ergänzen": "Figure 2.7 - Add screenshot",
    "8  CAN und Services prüfen": "8  Check CAN and services",
    "Das passende CAN1-Profil auswählen. Für neuere Microlino’s nach der Pioneer-Serie das Profil V2 wählen. Ohne CAN-Änderung ist kein Neustart nötig. History Cache und ABRP sind optional.": "Select the appropriate CAN1 profile. For newer Microlino vehicles produced after the Pioneer series, select the V2 profile. No restart is required if the CAN settings remain unchanged. History Cache and ABRP are optional.",
    "BITTE CAN2 NICHT ÄNDERN!": "DO NOT CHANGE CAN2!",
    "AWS immer aktiviert lassen wenn der Adapter mit dem MOT Portal genutzt werden soll.": "Always leave AWS enabled when the adapter is used with the MOT Portal.",
    "History cache erlaubt das lokale Speichern einiger Daten falls es einen Unterbruch zum Internet gibt. Die Daten werden bei Verfügbarkeit des Internets in das MOT Portal übertragen. GPS (Positionsdaten und Ladezustände werden nicht gespeichert).": "History Cache stores selected data locally during an internet outage and transfers it to the MOT Portal when internet access returns. Only SOC and vehicle speed are cached; GPS positions and other CAN data are not stored.",
    "ABRP nur aktivieren wenn Token und Key bekannt sind und der entsprechende Service bei ABRP abonniert ist.": "Enable ABRP only if the token and API key are available and the required ABRP service is subscribed.",
    "Abbildung 2.8 - CAN Einstellungen": "Figure 2.8 - CAN settings",
    "9  Validation starten": "9  Start validation",
    "Mit Start Validation die Geräte- und Telemetrie-Prüfung starten. Ist der Adapter bereits mit dem Fahrzeug verbunden, werden hier bereits Fahrzeugdaten angezeigt": "Select Start Validation to begin the device and telemetry checks. If the adapter is already connected to the vehicle, vehicle data is displayed here.",
    "Abbildung 2.9 - Validation: CAN Daten nur wenn der Adapter mit dem Fahrzeug verbunden ist": "Figure 2.9 - Validation: CAN data is available only when the adapter is connected to the vehicle",
    "10  Adapter bereit": "10  Adapter ready",
    "Der lokale Wizard ist abgeschlossen. Der Adapter ist für die nächste Phase und die Verbindung mit dem Fahrzeug beziehungsweise Portal bereit.": "The local wizard is complete. The adapter is ready for the next phase and for connection to the vehicle and portal.",
    "Hinweis: Nach dem Abschluss der Konfiguration stellt der Adapter kein lokales WLAN (MOT-xxxx) mehr bereit, ausser es wird keine Verbindung zu einem konfigurierten WLAN hergestellt. Der Adapter ist danach über die IP-Adresse im entsprechenden WLAN erreichbar. Wenn der Adapter mit dem MOT Portal verbunden und dem Benutzerkonto verknüpft ist (nächster Schritt) wird im Dashboard des Portals die IP-Adresse des Adapters angezeigt.": "Note: After setup is complete, the adapter no longer provides the local MOT-xxxx Wi-Fi network unless it cannot connect to a configured Wi-Fi network. You can then reach the adapter through its IP address on the connected network. After the adapter is linked to the MOT Portal and the user account in the next phase, its IP address is shown in the portal dashboard.",
    "Abbildung 2.10 -Completion": "Figure 2.10 - Completion",
    "Adapter und Portal verbinden": "Link the adapter and portal",
    "Jetzt kommt der Adapter ans Fahrzeug. Der Claim verbindet den bestätigten Benutzer nachvollziehbar mit der Vehicle-ID.\n\nHinweis: \nDer Adapter bracht immer Strom über den USB-C Anschluss. Es wird aktuell kein Strom über den OBD2 bezogen.": "The adapter is now connected to the vehicle. The claim securely links the confirmed user to the Vehicle ID. Note: The adapter always requires power through its USB-C port. It currently receives no power through OBD-II.",
    "1  Am Fahrzeug anschliessen": "1  Connect to the vehicle",
    "Hinweis: Der Adapter muss für diesen Schritt noch nicht zwingend mit dem Microlino aber mit dem WLAN und Internet verbunden sein.": "Note: For this step, the adapter does not yet have to be connected to the Microlino, but it must be connected to Wi-Fi and the internet.",
    "Adapter gemäss Einbauanleitung verbinden und sicherstellen, dass das konfigurierte WLAN erreichbar ist.\nAbbildung 3.1 - OBD2 Frontdoor": "Connect the adapter according to the installation instructions and make sure the configured Wi-Fi network is available. Figure 3.1 - OBD-II connector behind the front service flap",
    "Beispiel eines Mobil-WIFI Hotspot Stick:": "Example of a mobile Wi-Fi hotspot device:",
    "2  Erste Telemetrie abwarten": "2  Wait for the first telemetry",
    "Der Adapter wird bei AWS online und publiziert erstmals Daten für seine vorbereitete Vehicle-ID.": "The adapter comes online in AWS and publishes data for its prepared Vehicle ID for the first time.",
    "3  Claim bereitstellen (durch MOT Administrator)": "3  Provide the claim (MOT administrator)",
    "MOT prüft Gerät und erste Publikation, erstellt den Claim für die Vehicle-ID und gibt dem Benutzer den Claim bekannt.": "MOT verifies the device and its first publication, creates the claim for the Vehicle ID and provides the claim to the user.",
    "Abbildung 3.3 - Admin - Claim": "Figure 3.3 - Administrator - Claim",
    "4  Im Portal anmelden": "4  Sign in to the portal",
    "https://www.microlino-open-telemetry.ch/dashboard/ öffnen und mit dem persönlichen Passwort anmelden.": "Open https://www.microlino-open-telemetry.ch/dashboard/ and sign in using your personal password.",
    "Abbildung 3.4 - Vehicle Claim": "Figure 3.4 - Vehicle Claim",
    "5  Claim einlösen": "5  Redeem the claim",
    "Claim-Code im Portal eingeben und bestätigen. Er dient als Nachweis des Gerätezugriffs.": "Enter and confirm the claim code in the portal. It proves access to the device.",
    "Abbildung 3.5 - Vehicle Claim": "Figure 3.5 - Vehicle Claim",
    "6  Fahrzeug kontrollieren": "6  Check the vehicle",
    "Das Fahrzeug erscheint in der Auswahlliste. Live-Daten und Fahrzeugstatus prüfen.": "The vehicle appears in the vehicle selector. Check its live data and vehicle status.",
    "Abbildung 3.6 - Vehicle List": "Figure 3.6 - Vehicle list",
    "7  Optionale Dienste aktivieren": "7  Enable optional services",
    "Bei Bedarf über MOT freigeben: History, E-Mail-Benachrichtigung und SMS. E-Mail und SMS separat bestätigen; auch hier den Spam-Ordner prüfen.": "Ask MOT to enable optional services as required: History, email notifications and SMS. Confirm email and SMS separately, and check the spam folder if necessary.",
    "Abbildung 3.7 - Optional Service and Settings": "Figure 3.7 - Optional services and settings",
    "Arbeitsversion mit Screenshot-Platzhaltern": "Working version with screenshot placeholders",
    "Jeden grauen Platzhalter kannst du in Word auswählen und über Bild ändern durch den passenden Screenshot ersetzen.": "In Word, select each grey placeholder and use Change Picture to replace it with the appropriate screenshot.",
    "Passwort aufbewahren": "Store the password safely",
    "Dieses Passwort wird anschliessend sowohl für das lokale WLAN MOT-xxxx als auch für die Anmeldung mit admin benötigt.": "You will need this password both for the local MOT-xxxx Wi-Fi network and for signing in as admin.",
    "Falls das alte WLAN-Passwort gespeichert ist": "If the old Wi-Fi password is saved",
    "Computer und Smartphones fragen das neue Passwort möglicherweise erst nach bis zu 2 Minuten ab. Bei Bedarf WLAN kurz trennen oder das gespeicherte Netzwerk neu verbinden.": "Computers and smartphones may take up to two minutes before asking for the new password. If necessary, briefly disconnect Wi-Fi or reconnect to the saved network.",
    "Lokales WLAN bleibt verfügbar": "The local Wi-Fi network remains available",
    "Bis zum Abschluss der Inbetriebnahme bleibt das lokale WLAN MOT-xxxx aktiv - auch wenn Home- oder Mobile/WiFi2 bereits verbunden ist.": "The local MOT-xxxx Wi-Fi network remains active until commissioning is complete, even if Home Wi-Fi or Mobile/WiFi2 is already connected.",
    "Nächster Schritt": "Next step",
    "Adapter ans Fahrzeug anschliessen, erste AWS-Telemetrie abwarten und danach den Claim-Prozess durchführen.": "Connect the adapter to the vehicle, wait for the first AWS telemetry and then complete the claim process.",
    "Onboarding abgeschlossen": "Onboarding complete",
    "Fahrzeug ist im Portal sichtbar, Live-Daten sind geprüft und gewünschte optionale Dienste sind aktiviert.": "The vehicle is visible in the portal, live data has been checked and the requested optional services are enabled.",
}


def normalized(text: str) -> str:
    return text.strip()


def translate_paragraph(paragraph) -> bool:
    source = normalized(paragraph.text)
    if not source:
        return False
    target = TRANSLATIONS.get(source)
    if target is None:
        return False
    text_nodes = paragraph._p.xpath(".//w:t")
    if not text_nodes:
        paragraph.add_run(target)
    else:
        text_nodes[0].text = target
        for node in text_nodes[1:]:
            node.text = ""
    return True


def all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def main():
    document = Document(SOURCE)
    translated = 0
    missing = []
    for paragraph in all_paragraphs(document):
        source = normalized(paragraph.text)
        if not source:
            continue
        if translate_paragraph(paragraph):
            translated += 1
        elif source not in {"PHASE 1", "PHASE 2", "PHASE 3", "https://www.microlino-open-telemetry.ch/dashboard/"}:
            missing.append(source)

    if missing:
        raise RuntimeError("Missing translations:\n" + "\n".join(sorted(set(missing))))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "MOT Remote Pilot Onboarding - Extended Setup Guide"
    document.core_properties.subject = "English remote pilot onboarding guide"
    document.save(OUTPUT)
    print(f"translated={translated} output={OUTPUT}")


if __name__ == "__main__":
    main()
